from pathlib import Path

import docx

from ..base import BaseParser


class WordParser(BaseParser):
    """docx -> Markdown 转换逻辑 (增强版：支持表格与安全标题解析)。

    入参从 ``bytes`` 切换为 ``Path``：python-docx 原生支持基于路径打开，避免在内存中
    构造 ``BytesIO(file_stream)`` 副本。
    """

    def parse(self, source: Path | None) -> str:
        self.validate_source(source)
        if source is None:
            raise ValueError("WordParser 不支持 source=None 入参")
        doc = docx.Document(str(source))
        md_lines = []

        # 1. 解析普通段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name
            if style_name.startswith('Heading'):
                try:
                    level_str = style_name.replace('Heading', '').strip()
                    level = int(level_str) if level_str.isdigit() else 1
                    md_lines.append(f"{'#' * level} {text}")
                except Exception:
                    md_lines.append(f"# {text}")
            elif 'List Bullet' in style_name:
                md_lines.append(f"- {text}")
            elif 'List Number' in style_name:
                md_lines.append(f"1. {text}")
            else:
                md_lines.append(text)

        # 2. 提取表格内容
        if doc.tables:
            md_lines.append("\n### 文档表格数据\n")
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    md_lines.append(" | ".join(row_data))
                md_lines.append("\n")

        self.metadata['pages_or_length'] = len(doc.paragraphs)
        return "\n\n".join(md_lines)
