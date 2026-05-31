"""WordParser 端到端用例（mammoth → 复用 HTML 引擎）。

fixture 用 python-docx 程序化生成，覆盖 acceptance.feature 16 Scenario。
HTML 模块零回归由独立的 HTML 单测/集成承接（本模块不改 src/core/parser/html）。
"""

import io
import tempfile
from pathlib import Path

import docx
import pytest
from docx.oxml.ns import qn
from docx.shared import Inches

from src.core.exceptions import ParseBaseException
from src.core.parser.factory import ParserFactory
from src.core.parser.html import image_rewriter as image_rewriter_mod
from src.core.parser.providers.html_parser import HtmlParser
from src.core.parser.providers.pdf_parser import PdfParser
from src.core.parser.providers.word_parser import WordParser

_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06"
    b"\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05"
    b"\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = docx.oxml.shared.OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = docx.oxml.shared.OxmlElement("w:r")
    rpr = docx.oxml.shared.OxmlElement("w:rPr")
    run.append(rpr)
    t = docx.oxml.shared.OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


# parser 协议为 ``parse(source: Path)``：测试把字节落到临时文件再传路径。
_TMP_DIR = Path(tempfile.mkdtemp(prefix="word-parser-test-"))
_seq = 0


def _as_path(content: bytes) -> Path:
    global _seq
    _seq += 1
    path = _TMP_DIR / f"doc-{_seq}.docx"
    path.write_bytes(content)
    return path


def _docx_bytes(build) -> bytes:
    d = docx.Document()
    build(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _parse(build) -> tuple[str, dict]:
    parser = WordParser()
    md = parser.parse(_as_path(_docx_bytes(build)))
    return md, parser.extract_metadata()


# ==== 主流程 ====


def test_headings_keep_level_and_order():
    def build(d):
        d.add_heading("产品白皮书", 1)
        d.add_heading("一、概述", 2)
        d.add_heading("1.1 背景", 3)
        d.add_paragraph("正文内容占位，确保文档有有效内容。")

    md, _ = _parse(build)
    assert "# 产品白皮书" in md
    assert "## 一、概述" in md
    assert "### 1.1 背景" in md
    assert md.index("# 产品白皮书") < md.index("## 一、概述") < md.index("### 1.1 背景")


def test_paragraph_bold_and_hyperlink():
    def build(d):
        p = d.add_paragraph("本文介绍系统架构 ")
        p.add_run("重点说明").bold = True
        p2 = d.add_paragraph("详见 ")
        _add_hyperlink(p2, "https://example.com/docs", "官方文档")

    md, _ = _parse(build)
    assert "本文介绍系统架构" in md
    assert "**重点说明**" in md
    assert "[官方文档](https://example.com/docs)" in md


def test_nested_lists_preserve_hierarchy():
    def build(d):
        d.add_paragraph("第一点", style="List Bullet")
        d.add_paragraph("子项 A", style="List Bullet 2")
        d.add_paragraph("孙项 x", style="List Bullet 3")
        d.add_paragraph("第二点", style="List Bullet")
        d.add_paragraph("步骤一", style="List Number")
        d.add_paragraph("步骤一子步", style="List Number 2")

    md, _ = _parse(build)
    lines = md.splitlines()

    def indent(token: str) -> int:
        line = next(line for line in lines if token in line)
        return len(line) - len(line.lstrip(" "))

    # 真断层级：子项缩进 > 父项，孙项缩进 > 子项；同级父项缩进相等。
    assert "- 第一点" in md and "1. 步骤一" in md
    assert indent("子项 A") > indent("第一点")
    assert indent("孙项 x") > indent("子项 A")
    assert indent("第二点") == indent("第一点")
    assert indent("步骤一子步") > indent("步骤一")


def test_table_image_paragraph_order():
    def build(d):
        d.add_paragraph("表格前说明")
        t = d.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).text = "字段"
        t.cell(0, 1).text = "说明"
        t.cell(1, 0).text = "a"
        t.cell(1, 1).text = "甲"
        d.add_picture(io.BytesIO(_PNG), width=Inches(1))
        d.add_paragraph("表格后说明")

    md, _ = _parse(build)
    assert md.index("表格前说明") < md.index("| 字段 | 说明 |")
    assert md.index("| 字段 | 说明 |") < md.index("![](mock-minio://")
    assert md.index("![](mock-minio://") < md.index("表格后说明")


# ==== 表格 ====


def test_simple_table_is_markdown_table_in_place():
    def build(d):
        d.add_paragraph("说明段落")
        t = d.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).text = "字段"
        t.cell(0, 1).text = "说明"
        t.cell(1, 0).text = "user_id"
        t.cell(1, 1).text = "用户ID"
        t.cell(2, 0).text = "name"
        t.cell(2, 1).text = "姓名"

    md, meta = _parse(build)
    assert "| 字段 | 说明 |" in md
    assert "| --- | --- |" in md
    assert "| user_id | 用户ID |" in md
    assert "### 文档表格数据" not in md
    assert meta["table_count"] == 1


def test_merged_cell_table_is_record_style():
    def build(d):
        d.add_paragraph("正文占位")
        t = d.add_table(rows=3, cols=3)
        t.style = "Table Grid"
        for (r, c), v in {
            (0, 0): "模块",
            (0, 1): "接口",
            (0, 2): "说明",
            (1, 0): "用户",
            (2, 0): "用户",
            (1, 1): "创建",
            (1, 2): "新建用户",
            (2, 1): "删除",
            (2, 2): "移除用户",
        }.items():
            t.cell(r, c).text = v
        t.cell(1, 0).merge(t.cell(2, 0))

    md, meta = _parse(build)
    assert "[HTML表格开始：" in md
    assert "表格类型：记录式表格" in md
    assert "记录 1：" in md
    assert "<table" not in md
    assert meta["record_table_count"] >= 1


def test_single_table_failure_keeps_position(monkeypatch):
    # 真实 docx 难稳定触发表格处理异常，故让 HtmlTableProcessor 内部抛错，
    # 由其自带 try/except 转为原位失败兜底（失败兜底渲染本身由 HTML 模块既有测试守护）。
    from src.core.parser.html import renderer as renderer_mod

    def boom(self, table):
        raise RuntimeError("forced table failure")

    monkeypatch.setattr(renderer_mod.HtmlTableProcessor, "_classify_table", boom)

    def build(d):
        d.add_paragraph("失败前")
        t = d.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).text = "a"
        t.cell(0, 1).text = "b"
        t.cell(1, 0).text = "c"
        t.cell(1, 1).text = "d"
        d.add_paragraph("失败后")

    md, meta = _parse(build)
    assert "失败前" in md and "失败后" in md
    assert md.index("失败前") < md.index("失败后")
    assert meta["table_failure_count"] == 1


# ==== 图片 ====


def test_embedded_image_to_mock_minio():
    def build(d):
        d.add_paragraph("配图说明")
        d.add_picture(io.BytesIO(_PNG), width=Inches(1))

    md, meta = _parse(build)
    assert "![](mock-minio://" in md
    assert "data:image" not in md
    assert meta["image_count"] == 1


def test_image_object_path_failure_keeps_placeholder(monkeypatch):
    monkeypatch.setattr(
        image_rewriter_mod.HtmlImageRewriter,
        "build_mock_object_url",
        lambda self, url: None,
    )

    def build(d):
        d.add_paragraph("配图说明")
        d.add_picture(io.BytesIO(_PNG), width=Inches(1))

    md, meta = _parse(build)
    assert "mock-minio://unresolved/" in md
    assert meta["image_warning_count"] == 1


# ==== 异常与边界 ====


@pytest.mark.parametrize(
    "blob",
    [
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy-doc-ole",  # legacy .doc OLE 头
        b"PK\x03\x04corrupt-not-a-real-zip",  # 损坏的伪 zip
        b"\x00\x01\x02\x03not-ooxml-binary",  # 任意二进制
    ],
)
def test_non_ooxml_inputs_fail_fast(blob):
    with pytest.raises(ParseBaseException, match="非 .docx"):
        WordParser().parse(_as_path(blob))


def test_empty_stream_fails():
    with pytest.raises(ValueError, match="文件流不可为空"):
        WordParser().parse(_as_path(b""))


def test_mammoth_empty_content_fails():
    def build(d):
        d.add_paragraph("")  # 无有效正文

    with pytest.raises(ParseBaseException, match="无有效内容"):
        WordParser().parse(_as_path(_docx_bytes(build)))


def test_does_not_affect_non_word_parser_dispatch():
    assert isinstance(ParserFactory.get_parser("pdf"), PdfParser)
    assert isinstance(ParserFactory.get_parser("html"), HtmlParser)
    assert isinstance(ParserFactory.get_parser("docx"), WordParser)


def test_pipeline_contract_outputs_str_and_metadata():
    def build(d):
        d.add_heading("标题", 1)
        d.add_paragraph("正文内容占位足够长以产出有效 Markdown。")

    parser = WordParser()
    md = parser.parse(_as_path(_docx_bytes(build)))
    assert isinstance(md, str) and md.strip()
    meta = parser.extract_metadata()
    for key in ("table_count", "record_table_count", "table_failure_count", "image_count"):
        assert key in meta
